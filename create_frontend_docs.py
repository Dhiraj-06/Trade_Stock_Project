"""
Script to generate Frontend_API_Integration_Guide.docx Word Document for the Frontend Engineering Team.
"""
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from pathlib import Path
import numpy as np
import pandas as pd

BASE_DIR = Path(__file__).resolve().parent

def build_docx():
    doc = docx.Document()

    # Page Margins Setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Theme Color Palette
    COLOR_PRIMARY = RGBColor(15, 23, 42)    # Slate 900
    COLOR_SECONDARY = RGBColor(37, 99, 235) # Blue 600
    COLOR_TEXT = RGBColor(51, 65, 85)      # Slate 700

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('NIFTY 50 ML TRADING ENGINE')
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run('Frontend API Integration Guide & Complete Data Attributes Schema')
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = COLOR_SECONDARY
    run_sub.font.italic = True

    doc.add_paragraph()

    def add_heading_1(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return h

    def add_body_p(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
        run.font.color.rgb = COLOR_TEXT
        return p

    # Section 1
    add_heading_1('1. Executive Overview & API Endpoints')
    add_body_p('This integration document specifies all input parameters, endpoint paths, and JSON output data attributes provided by the backend ML Service for frontend dashboard integration.')

    table_urls = doc.add_table(rows=3, cols=3)
    table_urls.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Environment', 'Base API URL', 'Notes']
    for i, h in enumerate(headers):
        cell = table_urls.cell(0, i)
        cell.paragraphs[0].text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '1E293B')
        set_cell_margins(cell)

    data_urls = [
        ['Production (Cloud)', 'https://nifty-ml-backend.onrender.com', 'Live Deployed Render Backend with CORS enabled'],
        ['Local Development', 'http://localhost:8000', 'Local FastAPI dev server instance']
    ]
    for row_idx, row_data in enumerate(data_urls, start=1):
        bg_color = 'F8FAFC' if row_idx % 2 == 1 else 'FFFFFF'
        for col_idx, text in enumerate(row_data):
            cell = table_urls.cell(row_idx, col_idx)
            cell.paragraphs[0].text = text
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell)

    doc.add_paragraph()

    # Section 2
    add_heading_1('2. Input Parameters (Request Specification)')
    add_body_p('The frontend client can send request parameters to fetch real-time ML predictions and evaluate Groww custom limit orders.')

    table_inputs = doc.add_table(rows=4, cols=5)
    table_inputs.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_in = ['Parameter', 'Location', 'Type', 'Default', 'Description']
    for i, h in enumerate(headers_in):
        cell = table_inputs.cell(0, i)
        cell.paragraphs[0].text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '1E293B')
        set_cell_margins(cell)

    inputs_data = [
        ['ticker', 'Path Parameter', 'String', 'REQUIRED', 'NIFTY 50 stock ticker or company alias (e.g. BHARTIARTL, ITC, TCS, WIPRO, RELIANCE)'],
        ['qty', 'Query Parameter', 'Integer', '100', 'Position share quantity for calculating Required Capital, Profit, and Risk'],
        ['limit_price', 'Query Parameter', 'Float', 'null (Market LTP)', 'Trader proposed Limit Buy/Sell Price for evaluating Groww custom limit order']
    ]
    for row_idx, row_data in enumerate(inputs_data, start=1):
        bg_color = 'F8FAFC' if row_idx % 2 == 1 else 'FFFFFF'
        for col_idx, text in enumerate(row_data):
            cell = table_inputs.cell(row_idx, col_idx)
            cell.paragraphs[0].text = text
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell)

    doc.add_paragraph()

    # Section 3
    add_heading_1('3. Complete Backend JSON Output Attributes & UI Binding')
    add_body_p('Below is the complete reference of all JSON fields returned by GET /predict/{ticker} and how to bind them to UI components.')

    add_heading_2('A. Root Object Attributes')
    table_root = doc.add_table(rows=9, cols=4)
    table_root.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_root = ['JSON Attribute Key', 'Data Type', 'Sample Value', 'Description & UI Component Binding']
    for i, h in enumerate(headers_root):
        cell = table_root.cell(0, i)
        cell.paragraphs[0].text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '1E293B')
        set_cell_margins(cell)

    root_data = [
        ['ticker', 'String', '"BHARTIARTL"', 'Clean ticker symbol displayed in card headers'],
        ['current_price', 'Float', '1939.10', 'Real-Time Market Price (LTP) fetched directly from FYERS API'],
        ['predicted_price', 'Float', '1945.40', 'Full-Day EOD Expected Price target predicted by ML model'],
        ['predicted_return_pct', 'Float', '+0.32', 'Expected percentage return over full trading day (+0.32%)'],
        ['direction', 'String', '"UP" / "DOWN"', 'Primary trend direction predicted by model ("UP" = Call, "DOWN" = Put)'],
        ['proba_up', 'Float', '0.753', 'Raw probability score for UP move [0.0 to 1.0]'],
        ['confidence_score', 'Float', '0.506', 'Normalized model confidence score [0.0 to 1.0] (0.5 = 100% confidence)'],
        ['is_live_fyers', 'Boolean', 'true', 'Badge indicator: true = Live FYERS Feed, false = Offline Baseline']
    ]
    for row_idx, row_data in enumerate(root_data, start=1):
        bg_color = 'F8FAFC' if row_idx % 2 == 1 else 'FFFFFF'
        for col_idx, text in enumerate(row_data):
            cell = table_root.cell(row_idx, col_idx)
            cell.paragraphs[0].text = text
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell)

    doc.add_paragraph()

    add_heading_2('B. groww_order_analysis Object (Groww Custom Order AI Card)')
    table_groww = doc.add_table(rows=10, cols=4)
    table_groww.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers_root):
        cell = table_groww.cell(0, i)
        cell.paragraphs[0].text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '1E293B')
        set_cell_margins(cell)

    groww_data = [
        ['qty', 'Integer', '100', 'Position share quantity entered by user'],
        ['limit_price', 'Float', '1940.00', 'Trader limit price analyzed by AI'],
        ['required_capital', 'Float', '194000.00', 'Total capital required (Qty x Limit Price). Format: ₹1,94,000.00'],
        ['custom_profit_potential', 'Float', '630.00', '30-minute Rupee Profit Potential (+₹630.00)'],
        ['custom_max_risk', 'Float', '597.10', '30-minute Rupee Max Risk (-₹597.10)'],
        ['custom_rr_ratio', 'String', '"1:1.10"', 'Custom Order Risk/Reward Ratio string ("1:1.10")'],
        ['is_limit_in_entry_zone', 'Boolean', 'true', 'true if proposed limit price is inside optimal entry zone'],
        ['order_verdict', 'String', '"ORDER APPROVED"', 'AI Order Verdict Header (Green = APPROVED, Yellow = ADVISORY, Red = REJECTED)'],
        ['order_advice', 'String', '"Limit Price (₹1940.0)..."', 'Human-readable detailed AI advice text']
    ]
    for row_idx, row_data in enumerate(groww_data, start=1):
        bg_color = 'F8FAFC' if row_idx % 2 == 1 else 'FFFFFF'
        for col_idx, text in enumerate(row_data):
            cell = table_groww.cell(row_idx, col_idx)
            cell.paragraphs[0].text = text
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell)

    doc.add_paragraph()

    add_heading_2('C. ai_insights Object (Actionable Signal Hero & Technical Cards)')
    table_ai = doc.add_table(rows=10, cols=4)
    table_ai.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers_root):
        cell = table_ai.cell(0, i)
        cell.paragraphs[0].text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '1E293B')
        set_cell_margins(cell)

    ai_data = [
        ['actionable_signal', 'String', '"STRONG BUY"', 'HERO SIGNAL BOX BADGE (STRONG BUY, BUY, WAIT / POOR R:R, SELL)'],
        ['target_price_30m', 'Float', '1945.40', '30-Minute Intraday Target Price (₹1,945.40)'],
        ['target_return_30m_pct', 'Float', '+0.32', '30-Minute expected percentage return (+0.32%)'],
        ['suggested_entry_zone', 'String', '"₹1937.16 - ₹1948.80"', 'Fixed 30-Minute Optimal Entry Zone range string'],
        ['risk_reward_ratio', 'String', '"1:1.10"', 'Standard 30-Minute Risk/Reward Ratio string'],
        ['capital_allocation', 'String', '"100% Allocation"', 'Recommended Capital Position Sizing label'],
        ['confidence_pct', 'Float', '75.3', 'Model confidence score percentage (75.3%)'],
        ['market_trend', 'String', '"Strong Bullish"', 'Overall technical market trend description'],
        ['horizon_explanation', 'String', '"In next 30 mins..."', 'Natural language summary text paragraph']
    ]
    for row_idx, row_data in enumerate(ai_data, start=1):
        bg_color = 'F8FAFC' if row_idx % 2 == 1 else 'FFFFFF'
        for col_idx, text in enumerate(row_data):
            cell = table_ai.cell(row_idx, col_idx)
            cell.paragraphs[0].text = text
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell)

    doc.add_paragraph()

    # Section 4: Sample REST API Response JSON
    add_heading_1('4. Sample REST API Response Payload')
    add_body_p('Below is the exact JSON structure returned by GET /predict/BHARTIARTL?qty=100&limit_price=1940.00:')

    sample_json = '''{
  "ticker": "BHARTIARTL",
  "current_price": 1939.10,
  "predicted_price": 1945.40,
  "predicted_return_pct": 0.32,
  "direction": "UP",
  "proba_up": 0.753,
  "confidence_score": 0.506,
  "is_live_fyers": true,
  "groww_order_analysis": {
    "qty": 100,
    "limit_price": 1940.00,
    "required_capital": 194000.00,
    "custom_profit_potential": 630.00,
    "custom_max_risk": 597.10,
    "custom_rr_ratio": "1:1.10",
    "is_limit_in_entry_zone": true,
    "order_verdict": "🟢 ORDER APPROVED — EXCELLENT LIMIT ENTRY",
    "order_advice": "Limit Price (₹1940.0) is inside optimal Entry Zone (₹1937.16 - ₹1948.80) with a favorable Risk/Reward Ratio (1:1.10)."
  },
  "ai_insights": {
    "actionable_signal": "STRONG BUY",
    "target_price_30m": 1945.40,
    "target_return_30m_pct": 0.32,
    "target_price_eod": 1945.40,
    "target_return_eod_pct": 0.32,
    "full_day_rr_ratio": "1:1.50",
    "suggested_entry_zone": "₹1937.16 - ₹1948.80",
    "risk_reward_ratio": "1:1.10",
    "capital_allocation": "🟢 100% Capital Allocation (Full Position)",
    "confidence_pct": 75.3,
    "market_trend": "Strong Bullish",
    "horizon_explanation": "In the next 30 minutes, expected price target is ₹1945.40 (+0.32%)."
  },
  "risk_management": {
    "dynamic_stop_loss": 1933.13,
    "dynamic_target_price": 1945.40,
    "atr_14_points": 5.97,
    "risk_reward_ratio": "1:1.10",
    "passes_risk_reward_guard": true,
    "position_sizing": {
      "capital_allocation_pct": 100,
      "position_size_label": "🟢 100% Capital Allocation (Full Position)"
    },
    "key_levels_guard": {
      "support_20": 1937.16,
      "resistance_20": 1948.80,
      "near_resistance": false,
      "near_support": false,
      "suggested_entry_zone": "₹1937.16 - ₹1948.80"
    }
  }
}'''

    p_code = doc.add_paragraph()
    run_code = p_code.add_run(sample_json)
    run_code.font.name = 'Consolas'
    run_code.font.size = Pt(8.5)
    run_code.font.color.rgb = RGBColor(30, 41, 59)

    out_file = BASE_DIR.parent / "Frontend_API_Integration_Guide.docx"
    doc.save(out_file)
    print(f"SUCCESSFULLY SAVED DOCX TO: {out_file}")

if __name__ == "__main__":
    build_docx()
