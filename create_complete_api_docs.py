"""
Script to generate Complete_Frontend_API_Reference.docx Word Document.
Provides complete API reference, HTTP endpoints, JSON schemas, UI bindings, and JS integration code snippets for Frontend Developers.
"""
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent

def build_docx():
    doc = docx.Document()

    # Page Margins Setup
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Theme Color Palette
    COLOR_PRIMARY = RGBColor(15, 23, 42)    # Slate 900
    COLOR_SECONDARY = RGBColor(37, 99, 235) # Blue 600
    COLOR_TEXT = RGBColor(51, 65, 85)      # Slate 700

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('NIFTY 50 ML TRADING ENGINE')
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run('Complete Frontend API Reference, Attribute Schema & UI Mapping Guide')
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = COLOR_SECONDARY
    run_sub.font.italic = True

    doc.add_paragraph()

    def add_heading_1(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return h

    def add_body_p(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_TEXT
        return p

    def format_table(table, data, headers):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.cell(0, i)
            cell.paragraphs[0].text = h
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(cell, '1E293B')
            set_cell_margins(cell)

        for row_idx, row_data in enumerate(data, start=1):
            bg_color = 'F8FAFC' if row_idx % 2 == 1 else 'FFFFFF'
            for col_idx, text in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell.paragraphs[0].text = str(text)
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                set_cell_background(cell, bg_color)
                set_cell_margins(cell)

    # ----------------------------------------------------------------------
    # SECTION 1: Base URLs & Authentication
    # ----------------------------------------------------------------------
    add_heading_1('1. API Base URLs & Global Configuration')
    add_body_p('The backend ML Service is deployed on cloud infrastructure and provides CORS-enabled REST endpoints for frontend web/mobile applications.')

    url_headers = ['Environment', 'Base API URL', 'Description & Usage']
    url_data = [
        ['Production (Cloud)', 'https://nifty-ml-backend.onrender.com', 'Live cloud API server (CORS enabled for all origins)'],
        ['Local Development', 'http://localhost:8000', 'Local FastAPI dev server instance for testing']
    ]
    t_url = doc.add_table(rows=3, cols=3)
    format_table(t_url, url_data, url_headers)
    doc.add_paragraph()

    # ----------------------------------------------------------------------
    # SECTION 2: Complete API Endpoints Catalog
    # ----------------------------------------------------------------------
    add_heading_1('2. REST API Endpoints Catalog')
    add_body_p('Below is the list of all REST API endpoints exposed by the ML backend server:')

    api_cat_headers = ['Endpoint Path', 'HTTP Method', 'Purpose & Response Description']
    api_cat_data = [
        ['GET /predict/{ticker}', 'GET', 'Core Prediction & Order Analysis Engine (Returns ML predictions, Groww order analysis, risk guards, and technicals)'],
        ['GET /api/market/nifty50', 'GET', 'Returns live NIFTY 50 Index market summary (LTP, Day Change %, Market Hours status)'],
        ['GET /tickers', 'GET', 'Returns list of all 50 supported NIFTY 50 stock ticker symbols and company name aliases'],
        ['GET /health', 'GET', 'System health status, active champion model versions, and FYERS live connection status'],
        ['GET /fyers/login', 'GET', '1-Click OAuth2 login redirect endpoint for authorizing live FYERS API market data connection'],
        ['POST /fyers/token', 'POST', 'Directly accepts and persists FYERS access token JSON payload server-side']
    ]
    t_cat = doc.add_table(rows=7, cols=3)
    format_table(t_cat, api_cat_data, api_cat_headers)
    doc.add_paragraph()

    # ----------------------------------------------------------------------
    # SECTION 3: Endpoint 1 - GET /predict/{ticker} (Detailed Schema)
    # ----------------------------------------------------------------------
    add_heading_1('3. Endpoint Specification: GET /predict/{ticker}')
    add_body_p('This is the primary endpoint for the trading dashboard. It accepts a stock ticker symbol along with position quantity and custom limit price query parameters.')

    add_heading_2('Request Query & Path Parameters')
    req_headers = ['Parameter Name', 'Parameter Type', 'Data Type', 'Default', 'Description']
    req_data = [
        ['ticker', 'Path Parameter', 'String', 'REQUIRED', 'NIFTY 50 ticker symbol or alias (e.g. BHARTIARTL, ITC, TCS, WIPRO, RELIANCE)'],
        ['qty', 'Query Parameter', 'Integer', '100', 'Position share quantity for calculating Required Capital, Rupee Profit, and Rupee Risk'],
        ['limit_price', 'Query Parameter', 'Float', 'null (Market LTP)', 'Trader proposed Limit Entry Price for evaluating Groww custom limit order']
    ]
    t_req = doc.add_table(rows=4, cols=5)
    format_table(t_req, req_data, req_headers)
    doc.add_paragraph()

    add_heading_2('JSON Output Attribute Dictionary (Data Mapping Reference)')
    
    attr_headers = ['Attribute Path / Key', 'Data Type', 'Example Output', 'UI Component & Binding Description']
    attr_data = [
        ['ticker', 'String', '"BHARTIARTL"', 'Ticker symbol displayed in card titles and stock headers'],
        ['current_price', 'Float', '1939.10', 'Real-Time Market LTP fetched directly from FYERS API. Display as ₹1,939.10'],
        ['predicted_price', 'Float', '1945.40', 'Full-Day EOD Expected Price target predicted by ML model'],
        ['predicted_return_pct', 'Float', '+0.32', 'Expected percentage move (+0.32%). Color: Green for +, Red for -'],
        ['direction', 'String', '"UP" / "DOWN"', 'Trend direction ("UP" = Call/Buy, "DOWN" = Put/Sell)'],
        ['confidence_score', 'Float', '0.506', 'Normalized model confidence score [0.0 to 1.0]'],
        ['is_live_fyers', 'Boolean', 'true', 'Live Feed Badge: true = Live FYERS Feed, false = Offline Baseline'],
        ['groww_order_analysis.qty', 'Integer', '100', 'Position share quantity'],
        ['groww_order_analysis.limit_price', 'Float', '1940.00', 'Proposed limit entry price analyzed by AI'],
        ['groww_order_analysis.required_capital', 'Float', '194000.00', 'Total capital required (Qty x Limit Price). Format: ₹1,94,000.00'],
        ['groww_order_analysis.custom_profit_potential', 'Float', '630.00', '30-minute Rupee Profit Potential (+₹630.00)'],
        ['groww_order_analysis.custom_max_risk', 'Float', '597.10', '30-minute Rupee Max Risk (-₹597.10)'],
        ['groww_order_analysis.custom_rr_ratio', 'String', '"1:1.10"', 'Custom Order Risk/Reward Ratio string ("1:1.10")'],
        ['groww_order_analysis.is_limit_in_entry_zone', 'Boolean', 'true', 'true if limit price is inside optimal entry zone'],
        ['groww_order_analysis.order_verdict', 'String', '"🟢 ORDER APPROVED"', 'AI Order Verdict Header Badge (Green = APPROVED, Yellow = ADVISORY, Red = REJECTED)'],
        ['groww_order_analysis.order_advice', 'String', '"Limit Price (₹1940.0)..."', 'Human-readable detailed AI order guidance paragraph'],
        ['ai_insights.actionable_signal', 'String', '"STRONG BUY"', 'HERO SIGNAL BOX BADGE (STRONG BUY, BUY, WAIT / POOR R:R, SELL)'],
        ['ai_insights.target_price_30m', 'Float', '1945.40', '30-Minute Intraday Target Price (₹1,945.40)'],
        ['ai_insights.target_return_30m_pct', 'Float', '+0.32', '30-Minute expected percentage return (+0.32%)'],
        ['ai_insights.suggested_entry_zone', 'String', '"₹1937.16 - ₹1948.80"', 'Fixed 30-Minute Optimal Entry Zone range string'],
        ['ai_insights.risk_reward_ratio', 'String', '"1:1.10"', 'Standard 30-Minute Risk/Reward Ratio string'],
        ['ai_insights.capital_allocation', 'String', '"🟢 100% Allocation"', 'Recommended Capital Position Sizing label'],
        ['ai_insights.confidence_pct', 'Float', '75.3', 'Model confidence score percentage (75.3%)'],
        ['ai_insights.market_trend', 'String', '"Strong Bullish"', 'Overall technical market trend description'],
        ['ai_insights.horizon_explanation', 'String', '"In next 30 mins..."', 'Natural language summary text paragraph'],
        ['risk_management.dynamic_stop_loss', 'Float', '1933.13', '30-Minute Stop Loss level (Display in RED: ₹1,933.13)'],
        ['risk_management.dynamic_target_price', 'Float', '1945.40', '30-Minute Target level (Display in GREEN: ₹1,945.40)'],
        ['risk_management.atr_14_points', 'Float', '5.97', 'ATR Volatility in Points (₹5.97)'],
        ['risk_management.passes_risk_reward_guard', 'Boolean', 'true', 'Safety check: true if R:R >= 1.0'],
        ['analytics.momentum.rsi_14', 'Float', '58.42', 'RSI 14 momentum value. Highlight Overbought > 75 or Oversold < 25'],
        ['analytics.momentum.macd_status', 'String', '"Positive (Buying)"', 'MACD Crossover status description'],
        ['analytics.trade_score', 'Integer', '82', 'Overall AI Quality Trade Score out of 100 (Gauge bar)']
    ]
    t_attr = doc.add_table(rows=len(attr_data)+1, cols=4)
    format_table(t_attr, attr_data, attr_headers)
    doc.add_paragraph()

    # ----------------------------------------------------------------------
    # SECTION 4: Endpoints 2 to 6
    # ----------------------------------------------------------------------
    add_heading_1('4. Additional Supporting API Endpoints')

    add_heading_2('A. GET /api/market/nifty50')
    add_body_p('Returns live NIFTY 50 Index market summary for dashboard ticker ticker headers:')
    mkt_headers = ['Field Key', 'Data Type', 'Sample Output', 'Description']
    mkt_data = [
        ['symbol', 'String', '"NSE:NIFTY50-INDEX"', 'NIFTY 50 Index symbol'],
        ['price', 'Float', '24570.65', 'Live NIFTY 50 Index price'],
        ['change_percent', 'Float', '+0.45', 'NIFTY 50 Day change percentage'],
        ['market_status', 'String', '"OPEN" / "CLOSED"', 'Market hours status (OPEN / CLOSED)'],
        ['is_live', 'Boolean', 'true', 'Live market feed connection flag']
    ]
    t_mkt = doc.add_table(rows=6, cols=4)
    format_table(t_mkt, mkt_data, mkt_headers)
    doc.add_paragraph()

    add_heading_2('B. GET /tickers')
    add_body_p('Returns JSON array of all 50 supported NIFTY 50 tickers for populating search dropdowns:')
    add_body_p('Response: ["ADANIENT", "ADANIPORTS", "ASIANPAINT", "AXISBANK", "BAJFINANCE", "BHARTIARTL", "CIPLA", "COALINDIA", "HDFCBANK", "INFY", "ITC", "JSWSTEEL", "RELIANCE", "SBIN", "TCS", "WIPRO", ...]')

    add_heading_2('C. GET /health')
    add_body_p('System health check endpoint returning model version metadata and FYERS auth status.')

    doc.add_paragraph()

    # ----------------------------------------------------------------------
    # SECTION 5: Frontend JavaScript Integration Example
    # ----------------------------------------------------------------------
    add_heading_1('5. Frontend JavaScript Integration Example Code')
    add_body_p('Below is the standard JavaScript fetch() implementation for frontend developers:')

    js_code = '''// Standard JavaScript API Fetch Implementation
async function fetchStockPrediction(tickerSymbol, quantity = 100, limitPrice = null) {
    const baseUrl = 'https://nifty-ml-backend.onrender.com';
    let url = `${baseUrl}/predict/${encodeURIComponent(tickerSymbol)}?qty=${quantity}`;
    if (limitPrice !== null && limitPrice > 0) {
        url += `&limit_price=${limitPrice}`;
    }

    try {
        const response = await fetch(url, { method: 'GET', headers: { 'Accept': 'application/json' } });
        if (!response.ok) {
            const errData = await response.json();
            throw new Error(errData.detail || 'API request failed');
        }
        const data = await response.json();

        // 1. Bind Hero Signal
        document.getElementById('actionableSignal').innerText = data.ai_insights.actionable_signal;

        // 2. Bind Groww Order Evaluator Card
        const gr = data.groww_order_analysis;
        document.getElementById('orderVerdict').innerText = gr.order_verdict;
        document.getElementById('requiredCapital').innerText = '₹' + gr.required_capital.toLocaleString('en-IN');
        document.getElementById('customRR').innerText = gr.custom_rr_ratio;
        document.getElementById('customProfit').innerText = '+₹' + gr.custom_profit_potential.toFixed(2);
        document.getElementById('customRisk').innerText = '-₹' + gr.custom_max_risk.toFixed(2);
        document.getElementById('orderAdvice').innerText = gr.order_advice;

        // 3. Bind Risk Management Card
        const rm = data.risk_management;
        document.getElementById('stopLoss').innerText = '₹' + rm.dynamic_stop_loss.toFixed(2);
        document.getElementById('targetPrice').innerText = '₹' + rm.dynamic_target_price.toFixed(2);
        document.getElementById('entryZone').innerText = rm.key_levels_guard.suggested_entry_zone;

        return data;
    } catch (err) {
        console.error('Error fetching prediction:', err);
        alert('API Error: ' + err.message);
    }
}'''

    p_code = doc.add_paragraph()
    run_code = p_code.add_run(js_code)
    run_code.font.name = 'Consolas'
    run_code.font.size = Pt(8.5)
    run_code.font.color.rgb = RGBColor(30, 41, 59)

    out_file = BASE_DIR.parent / "Complete_Frontend_API_Reference.docx"
    doc.save(out_file)
    print(f"SUCCESSFULLY SAVED DOCX TO: {out_file}")

if __name__ == "__main__":
    build_docx()
