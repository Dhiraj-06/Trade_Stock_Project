"""
Script to generate Frontend_Page_Layout_Guide.docx Word Document for Frontend Dashboard Architecture.
"""
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent

def build_docx():
    doc = docx.Document()

    # Page Margins Setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Theme Color Palette
    COLOR_PRIMARY = RGBColor(15, 23, 42)    # Slate 900
    COLOR_SECONDARY = RGBColor(37, 99, 235) # Blue 600
    COLOR_TEXT = RGBColor(51, 65, 85)      # Slate 700

    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('NIFTY 50 ML TRADING ENGINE')
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run('Frontend Dashboard Page-by-Page Feature Mapping & Component Architecture')
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = COLOR_SECONDARY
    run_sub.font.italic = True

    doc.add_paragraph()

    def add_heading_1(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return h

    def add_body_p(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
        run.font.color.rgb = COLOR_TEXT
        return p

    # Overview
    add_heading_1('1. Dashboard Structure & Navigation Overview')
    add_body_p('The frontend trading platform is structured into 3 specialized dashboard pages to provide traders with real-time market insights, AI trade discovery, and risk management protection.')

    table_pages = doc.add_table(rows=4, cols=3)
    table_pages.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Page Name', 'Primary Purpose / UX Goal', 'Key Highlights']
    for i, h in enumerate(headers):
        cell = table_pages.cell(0, i)
        cell.paragraphs[0].text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '1E293B')
        set_cell_margins(cell)

    pages_summary = [
        ['1. Market Data Page', 'Real-Time Price Scanning & Market Overview', 'Live FYERS Feed, Ticker Search, Trend Badges, RSI/MACD Technicals'],
        ['2. Trade Discovery Page', 'ML Signal Generation & Groww Order AI Evaluator', 'Actionable Signal Hero Box, Custom Limit Order Evaluator, Target Horizons'],
        ['3. Risk Management Page', 'Capital Preservation & Trade Risk Control', 'Dynamic ATR Stop Loss, Position Sizing Guards, Fixed Entry Zones, Trade Score']
    ]
    for row_idx, row_data in enumerate(pages_summary, start=1):
        bg_color = 'F8FAFC' if row_idx % 2 == 1 else 'FFFFFF'
        for col_idx, text in enumerate(row_data):
            cell = table_pages.cell(row_idx, col_idx)
            cell.paragraphs[0].text = text
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell)

    doc.add_paragraph()

    # PAGE 1
    add_heading_1('2. PAGE 1: Market Data Page (Real-Time Overview & Scanner)')
    add_body_p('Goal: Provide real-time price monitoring, market connectivity status, technical momentum, and ticker selection.')

    add_heading_2('UI Components & Backend Attribute Bindings')
    table_p1 = doc.add_table(rows=6, cols=4)
    table_p1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_col = ['UI Section / Card Name', 'Displayed Features', 'Backend JSON Attribute Binding', 'Visual Rules & Styling']
    for i, h in enumerate(headers_col):
        cell = table_p1.cell(0, i)
        cell.paragraphs[0].text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '1E293B')
        set_cell_margins(cell)

    p1_data = [
        ['Header Bar', 'Live Connection Health Badge & FYERS OAuth Login Button', 'is_live_fyers (Boolean)', 'Green badge for FYERS Live API, Yellow badge for Offline Baseline. Button calls GET /fyers/login'],
        ['Ticker Scanner', 'NIFTY 50 Dropdown & Search Bar', 'GET /tickers API endpoint', 'Auto-complete select dropdown containing all 50 valid NIFTY 50 stock symbols'],
        ['Live Market Price Hero', 'Real-Time Price (LTP), Full-Day Target, Return %', 'current_price, predicted_price, predicted_return_pct', 'Large bold LTP display (e.g. ₹1,939.10). Green for positive return, Red for negative'],
        ['Technical Momentum', 'RSI 14 Value, MACD Signal Status', 'analytics.momentum.rsi_14, analytics.momentum.macd_status', 'RSI gauge bar. Highlight Overbought > 75 (Warning) or Oversold < 25 (Opportunity)'],
        ['Market Trend & Volatility', 'Market Trend Badge & 20-Candle Volatility %', 'ai_insights.market_trend, analytics.volatility_20_pct', 'Trend Badges: Strong Bullish / Bullish / Bearish / Strong Bearish']
    ]
    for row_idx, row_data in enumerate(p1_data, start=1):
        bg_color = 'F8FAFC' if row_idx % 2 == 1 else 'FFFFFF'
        for col_idx, text in enumerate(row_data):
            cell = table_p1.cell(row_idx, col_idx)
            cell.paragraphs[0].text = text
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell)

    doc.add_paragraph()

    # PAGE 2
    add_heading_1('3. PAGE 2: Trade Discovery Page (ML Signals & Order AI Evaluator)')
    add_body_p('Goal: Deliver actionable AI trade signals, evaluate trader custom limit orders, and display multi-horizon price targets.')

    add_heading_2('UI Components & Backend Attribute Bindings')
    table_p2 = doc.add_table(rows=5, cols=4)
    table_p2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers_col):
        cell = table_p2.cell(0, i)
        cell.paragraphs[0].text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '1E293B')
        set_cell_margins(cell)

    p2_data = [
        ['Actionable Signal Hero Box', 'Hero Signal Badge, Model Confidence %, Summary Text', 'ai_insights.actionable_signal, ai_insights.confidence_pct, horizon_explanation', 'Prominent Banner: Green for STRONG BUY / BUY, Yellow for WAIT / POOR R:R, Red for SELL'],
        ['Groww Order Input Terminal', 'Input Form: Quantity (Qty) & Limit Price (₹)', 'Query params: ?qty={qty}&limit_price={price}', 'Interactive form with input validation and "Analyze Groww Order" CTA button'],
        ['Groww Order AI Card', 'Required Capital, Verdict Banner, Custom R:R, Profit (+₹), Risk (-₹), Advice', 'groww_order_analysis object (required_capital, order_verdict, custom_profit_potential, custom_max_risk, order_advice)', 'Card Banner: Green border for ORDER APPROVED, Yellow for ORDER ADVISORY, Red for ORDER REJECTED'],
        ['Dual Target Horizons', '30-Min Intraday Target vs. Full-Day EOD Target', 'ai_insights.target_price_30m, target_return_30m_pct, target_price_eod, target_return_eod_pct', 'Comparison card showing 30-min expected move vs. full-day EOD move']
    ]
    for row_idx, row_data in enumerate(p2_data, start=1):
        bg_color = 'F8FAFC' if row_idx % 2 == 1 else 'FFFFFF'
        for col_idx, text in enumerate(row_data):
            cell = table_p2.cell(row_idx, col_idx)
            cell.paragraphs[0].text = text
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell)

    doc.add_paragraph()

    # PAGE 3
    add_heading_1('4. PAGE 3: Risk Management Page (Capital Protection & Entry Guards)')
    add_body_p('Goal: Protect trader capital using dynamic ATR stop losses, position sizing rules, key levels guards, and trade quality scores.')

    add_heading_2('UI Components & Backend Attribute Bindings')
    table_p3 = doc.add_table(rows=5, cols=4)
    table_p3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers_col):
        cell = table_p3.cell(0, i)
        cell.paragraphs[0].text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(cell, '1E293B')
        set_cell_margins(cell)

    p3_data = [
        ['Risk Management Dashboard Card', 'Dynamic Stop Loss, 30-Min Target, ATR Points, Risk/Reward Ratio', 'risk_management.dynamic_stop_loss, dynamic_target_price, atr_14_points, risk_reward_ratio', 'Highlight Stop Loss in RED (e.g. ₹1,933.13), Target in GREEN (e.g. ₹1,945.40), R:R formatted as "1:1.10"'],
        ['Capital Position Sizing Card', 'Position Size Allocation Label & Risk Rating Badge', 'risk_management.position_sizing.position_size_label, analytics.risk_rating', 'Allocation Badge: Green for 100% Full Position, Yellow for 50% Reduced Position, Red for 0% Wait'],
        ['Key Levels Guard Card', 'Fixed 30-Min Optimal Entry Zone, Support & Resistance', 'key_levels_guard.suggested_entry_zone, support_20, resistance_20', 'Display fixed Entry Zone range (e.g. ₹1,937.16 - ₹1,948.80). Highlight Support & Resistance levels'],
        ['AI Quality Trade Score', 'Trade Score Progress Meter (10 to 99)', 'analytics.trade_score (Integer)', 'Gauge/Progress Bar out of 100. Score > 75 = Excellent Setup, < 50 = Caution']
    ]
    for row_idx, row_data in enumerate(p3_data, start=1):
        bg_color = 'F8FAFC' if row_idx % 2 == 1 else 'FFFFFF'
        for col_idx, text in enumerate(row_data):
            cell = table_p3.cell(row_idx, col_idx)
            cell.paragraphs[0].text = text
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            set_cell_background(cell, bg_color)
            set_cell_margins(cell)

    out_file = BASE_DIR.parent / "Frontend_Page_Layout_Guide.docx"
    doc.save(out_file)
    print(f"SUCCESSFULLY SAVED DOCX TO: {out_file}")

if __name__ == "__main__":
    build_docx()
