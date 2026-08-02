"""
Script to generate NIFTY 50 ML Trading Model Project Report Word Document (.docx).
"""
import sys
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

BASE_DIR = Path(__file__).resolve().parent.parent

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()
    
    # Page setup - Margins 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles & Colors
    # Primary: Deep Navy (#0F172A), Accent: Cyan (#0284C7), Dark Text: #1E293B
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("NIFTY 50 ML TRADING MODEL ENGINE")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Fyers Live Integration, Scale-Invariant Features & Automated Self-Retraining\nFull Technical Project Documentation")
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
        return h

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        r2 = p.add_run(text)
        return p

    # 1. Executive Summary
    add_heading_1("1. Executive Summary & Project Goals")
    p = doc.add_paragraph("This document presents the completed technical implementation of the NIFTY 50 Machine Learning Trading Engine. Designed for live trading signal generation, the service processes historical and real-time market data across 50 NIFTY constituent stocks, engineers scale-invariant feature vectors, predicts 30-minute intraday returns and directions, serves predictions via REST API, and retrains itself unattended with strict champion promotion gating.")
    
    add_bullet("Project Scope:", "ML Core pipeline — data prep, feature engineering, walk-forward training, live Fyers API v3 integration, FastAPI endpoint serving, scheduled self-retraining, and drift monitoring.")
    add_bullet("Trading Horizon:", "Configured for 30-minute intraday prediction windows (2 x 15-minute candles).")
    add_bullet("Target Instruments:", "NIFTY 50 constituent universe (WIPRO, RELIANCE, TCS, TATASTEEL, EICHERMOT, etc.).")

    # 2. Folder Structure & Architecture
    add_heading_1("2. Folder Structure & Modular Architecture")
    doc.add_paragraph("The service is structured in a modular, decoupled architecture where each component has an explicit responsibility:")

    struct_text = """ml_service/
├── config/             # Central YAML settings loader & environment variables
├── data/               # Raw historical CSVs, processed datasets, and live Fyers buffers
├── features/           # Scale-invariant feature engineering (single source of truth)
├── training/           # Data prep, 5-fold walk-forward validation, regressor & classifier trainers
├── models/             # Versioned model registry (load_champion, promote, rollback)
├── inference/          # Fyers client wrapper, ticker validation, live pipeline & predictor
├── retraining/         # Retrain job, APScheduler cron runner, and drift monitor
├── api/                # FastAPI endpoints (/predict, /health, /retrain, /fyers/login)
├── scripts/            # CLI runners (run_training.py)
├── tests/              # Pytest unit & integration test suite (6/6 passing)
├── Dockerfile          # Production containerization
└── requirements.txt    # Production dependencies"""

    p_code = doc.add_paragraph()
    r_code = p_code.add_run(struct_text)
    r_code.font.name = 'Consolas'
    r_code.font.size = Pt(9.5)
    r_code.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    # 3. Core Technical Decisions
    add_heading_1("3. Core Machine Learning & Technical Decisions")
    
    add_heading_2("3.1 Raw OHLCV Ingestion & Indicator Cleanse")
    doc.add_paragraph("Pre-built indicator columns in the raw historical dataset were discarded to keep only raw OHLCV columns (Ticker, Date, Open, High, Low, Close, Volume). This ensures that historical training and live Fyers inference compute technical indicators using the exact same formulas, eliminating train/serve formula skew.")

    add_heading_2("3.2 Scale-Invariant Feature Engineering (~46 Features)")
    doc.add_paragraph("All price-scale and volume-scale columns are converted into scale-invariant ratios, percentage differences, or naturally bounded oscillators. This enables a single shared model to generalize across stocks priced from ₹160 to ₹7,800+:")
    add_bullet("Normalized Ratios:", "(ema_short - close) / close, (ema_long - close) / close, (vwap - close) / close, atr / close.")
    add_bullet("Bollinger Ratios:", "(bb_upper - close) / close, (close - bb_lower) / close, bb_width, bb_pct_b.")
    add_bullet("MACD & Lags:", "macd / close, macd_signal / close, close_lag_k_return, return_lag_k.")
    add_bullet("Bounded Oscillators:", "RSI (0-100), ADX (0-100), CCI, Stochastic %K/%D (0-100).")
    add_bullet("Per-Ticker Grouping:", "All rolling windows are strictly computed per ticker group to prevent cross-ticker boundary contamination.")

    add_heading_2("3.3 Zero-Leakage Walk-Forward Validation")
    doc.add_paragraph("Random k-fold split is strictly forbidden for time series data. Validation is performed using 5-fold TimeSeriesSplit (expanding window), ensuring the model never sees future data during training folds.")

    # 4. Model Performance & Metrics Table
    add_heading_1("4. Model Architecture & Trained Metrics")
    doc.add_paragraph("The engine utilizes a two-model architecture operating in parallel:")
    add_bullet("Model B — Regressor (XGBoost):", "Predicts expected 30-minute percentage return (target_return_pct).")
    add_bullet("Model C — Classifier (XGBoost):", "Predicts next 30-minute direction probability P(UP) / P(DOWN) and confidence score.")

    doc.add_paragraph("The models were trained on 41,444 dataset rows across 50 tickers. Below are the 5-fold walk-forward validation metrics for active registry champions:")

    # Table of Metrics
    table = doc.add_table(rows=9, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Model", "Metric Name", "Walk-Forward 5-Fold Champion Score"]
    for idx, h in enumerate(headers):
        cell = table.cell(0, idx)
        cell.paragraphs[0].text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)

    rows_data = [
        ("Model B (Regressor)", "Avg MAE (Mean Absolute Error)", "1.51399%"),
        ("Model B (Regressor)", "Avg RMSE (Root Mean Sq Error)", "2.02870%"),
        ("Model B (Regressor)", "Directional Accuracy", "54.56%"),
        ("Model C (Classifier)", "Avg Accuracy", "53.96%"),
        ("Model C (Classifier)", "Avg Precision", "54.23%"),
        ("Model C (Classifier)", "Avg Recall", "70.27%"),
        ("Model C (Classifier)", "Avg F1-Score", "0.6115"),
        ("Model C (Classifier)", "Avg ROC-AUC", "0.5545"),
    ]

    for row_idx, data in enumerate(rows_data, start=1):
        for col_idx, text in enumerate(data):
            cell = table.cell(row_idx, col_idx)
            cell.paragraphs[0].text = text
            if row_idx % 2 == 0:
                set_cell_background(cell, "F1F5F9")
            else:
                set_cell_background(cell, "FFFFFF")
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 5. Fyers API Integration & Auth Flow
    add_heading_1("5. Fyers API v3 Integration & Auth Flow")
    doc.add_paragraph("The service integrates natively with the Fyers API v3 SDK for live quote fetching and historical candle updates:")
    add_bullet("Credentials Management:", "FYERS_APP_ID and FYERS_SECRET_KEY stored securely in ml_service/.env.")
    add_bullet("1-Click OAuth Auth Code Flow:", "Exposes GET /fyers/login and GET /fyers/callback endpoints. The user logs in once, and the server automatically exchanges the auth code for FYERS_ACCESS_TOKEN and saves it to .env.")
    add_bullet("Rolling Buffer Maintainer:", "Maintains rolling intraday Parquet candle buffers in data/live/{ticker}_live.parquet to compute warm-up feature vectors.")
    add_bullet("Fallback Mode:", "If offline or unauthenticated, automatically falls back to dry-run mode without crashing.")

    # 6. Ticker Validation & Alias Resolution
    add_heading_1("6. Ticker Validation & Alias Resolution")
    doc.add_paragraph("To ensure input safety, inference/ticker_utils.py strictly enforces constituent boundaries:")
    add_bullet("Alias Resolution:", "Automatically maps user input variations like 'Tata Steel' or 'tata steel' -> TATASTEEL, 'Infosys' -> INFY, 'HDFC Bank' -> HDFCBANK, 'Bajaj Auto' -> BAJAJ-AUTO.")
    add_bullet("NIFTY 50 Boundary Guard:", "Non-NIFTY 50 inputs (like ZOMATO, PAYTM, AAPL) are rejected immediately with HTTP 400 Bad Request detailing that the model is trained exclusively on NIFTY 50 stocks.")

    # 7. Self-Retraining & Champion Promotion Gate
    add_heading_1("7. Automated Retraining & Champion Promotion Gate")
    doc.add_paragraph("The model requires zero manual retraining after deployment:")
    add_bullet("Cron Scheduler (retraining/scheduler.py):", "Runs automatically via APScheduler every weekday at 6:00 PM after Indian market close.")
    add_bullet("Champion Promotion Gate (retraining/retrain_job.py):", "Retrains candidate models on updated data, computes 5-fold walk-forward metrics, and compares against current champion metadata. Promotes candidate ONLY if metrics meet or exceed the champion.")
    add_bullet("Drift Monitoring (retraining/drift_monitor.py):", "Tracks live predictions vs realized outcomes. Triggers emergency retrain if directional accuracy drops below 52%.")

    # 8. Test Suite
    add_heading_1("8. Pytest Test Suite Results")
    doc.add_paragraph("All 6 unit and integration tests pass cleanly in 4.26 seconds:")
    add_bullet("test_no_cross_ticker_boundary_contamination:", "PASSED (Zero boundary leakage).")
    add_bullet("test_features_are_scale_invariant:", "PASSED (Scale-invariant feature ranges).")
    add_bullet("test_end_to_end_pipeline:", "PASSED (End-to-end data prep -> train -> predict).")
    add_bullet("test_valid_nifty50_ticker:", "PASSED (Canonical ticker verification).")
    add_bullet("test_alias_resolution:", "PASSED ('Tata Steel' -> TATASTEEL).")
    add_bullet("test_invalid_non_nifty50_ticker:", "PASSED ('ZOMATO' -> HTTP 400 rejection).")

    # 9. Summary & How to Run
    add_heading_1("9. Service Execution & Deployment")
    doc.add_paragraph("To run the live service and test predictions:")
    add_bullet("Command:", "cd ml_service && python -m uvicorn api.app:app --host 127.0.0.1 --port 8000 --reload")
    add_bullet("Web Dashboard:", "http://127.0.0.1:8000/")
    add_bullet("Docker Command:", "docker build -t nifty50-ml-service . && docker run -d -p 8000:8000 --env-file .env nifty50-ml-service")

    # Save Word Document
    out_file = BASE_DIR.parent / "NIFTY50_ML_Model_Project_Report.docx"
    doc.save(out_file)
    print("Project report successfully generated:", out_file)
    return out_file

if __name__ == "__main__":
    create_report()
