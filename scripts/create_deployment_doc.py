"""
Script to generate NIFTY 50 ML Trading Engine Deployment Guide Word Document (.docx).
"""
import sys
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

BASE_DIR = Path(__file__).resolve().parent.parent

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_deployment_doc():
    doc = docx.Document()
    
    # Page setup - Margins 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styles & Colors
    # Primary: Deep Navy (#0F172A), Accent: Cyan (#0284C7), Dark Text: #1E293B
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("NIFTY 50 ML TRADING ENGINE")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Production Deployment & Integration Guide\nStep-by-Step Instructions for ML Backend, Frontend & Fyers Live API")
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x02, 0x84, 0xC7)
        return h

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True
        r1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        r2 = p.add_run(text)
        return p

    def add_code_block(code_text):
        p = doc.add_paragraph()
        r = p.add_run(code_text)
        r.font.name = 'Consolas'
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        return p

    # 1. Executive Summary & Architecture Overview
    add_heading_1("1. Executive Summary & Deployment Architecture")
    doc.add_paragraph("This guide provides complete, step-by-step instructions for deploying the NIFTY 50 Machine Learning Trading Engine and connecting it to a production frontend web application. The platform follows a modern decoupled microservices architecture where the ML engine runs as a dedicated Python FastAPI service and communicates asynchronously with the web frontend via REST API endpoints.")

    arch_diagram = """
┌────────────────────────────────────────┐          HTTPS REST API Calls         ┌────────────────────────────────────────┐
│            Frontend Website            │ ────────────────────────────────────> │           ML Backend Engine            │
│      (Vercel / Netlify / Cloud)        │ <──────────────────────────────────── │  (AWS EC2 / DigitalOcean Cloud VM)     │
│   Renders UI Cards, Gauges & Charts    │        Returns Prediction JSON        │   FastAPI + Uvicorn + Systemd / Docker │
└────────────────────────────────────────┘                                       └────────────────────────────────────────┘
                                                                                                    ▲
                                                                                                    │ Live Intraday Data
                                                                                                    ▼
                                                                                 ┌────────────────────────────────────────┐
                                                                                 │          Fyers Live API v3             │
                                                                                 │     Quotes & 15m/30m Candles           │
                                                                                 └────────────────────────────────────────┘"""
    add_code_block(arch_diagram)

    # 2. Prerequisites
    add_heading_1("2. Deployment Prerequisites")
    doc.add_paragraph("Before launching the deployment, ensure you have gathered the following resources:")
    add_bullet("Server Instance:", "Cloud Virtual Machine (AWS EC2 t3.medium, DigitalOcean Droplet, or Hetzner VM) with Ubuntu 22.04 LTS, 2 vCPUs, 4GB RAM, 25GB SSD storage.")
    add_bullet("Domain Name:", "Registered domain or subdomain pointing to your Cloud VM IP address (e.g., api.yourdomain.com for backend, app.yourdomain.com for frontend).")
    add_bullet("Fyers API Credentials:", "Active Fyers API App ID (FYERS_APP_ID) and Secret Key (FYERS_SECRET_KEY) from fyers.in app dashboard.")
    add_bullet("Repository Access:", "Access to Git repository containing the ml_service codebase and frontend web code.")

    # 3. Deploying ML Backend Engine
    add_heading_1("3. Deploying the ML Backend Engine")
    doc.add_paragraph("Follow these exact commands on your Ubuntu Cloud VM server to install dependencies, train initial champions, and configure 24/7 background execution.")

    add_heading_2("Step 3.1: Environment Setup & Code Checkout")
    add_code_block("""# SSH into server
ssh ubuntu@YOUR_SERVER_IP

# Install system dependencies & python3-venv
sudo apt update && sudo apt install -y python3-pip python3-venv git curl

# Clone repository and navigate to ml_service
git clone https://github.com/your-org/ML_Model.git
cd ML_Model/ml_service

# Create and activate python virtual environment
python3 -m venv venv
source venv/bin/activate
pip install --upgrade pip
pip install -r requirements.txt""")

    add_heading_2("Step 3.2: Configure Environment Variables (.env)")
    doc.add_paragraph("Create the ml_service/.env configuration file on your server:")
    add_code_block("""nano .env

# Add your Fyers API credentials inside ml_service/.env:
FYERS_APP_ID=YOUR_APP_ID-100
FYERS_SECRET_KEY=YOUR_SECRET_KEY
FYERS_ACCESS_TOKEN=YOUR_24H_ACCESS_TOKEN""")

    add_heading_2("Step 3.3: Initial Model Training & Registry Verification")
    doc.add_paragraph("Run the one-shot training script to generate baseline 30-minute intraday champion models:")
    add_code_block("""# Run training script (trains regressor & classifier, populates models/registry/)
python scripts/run_training.py

# Verify that champion models are generated in registry
ls -la models/registry/return_regressor/champion
ls -la models/registry/direction_classifier/champion""")

    add_heading_2("Step 3.4: Configure Systemd Persistent Background Service")
    doc.add_paragraph("Create a Linux systemd service file to keep the ML FastAPI service running 24/7 and automatically restart on reboot:")
    add_code_block("""sudo nano /etc/systemd/system/ml_service.service

# Copy and paste the following service definition:
[Unit]
Description=NIFTY 50 ML Trading Engine & Self-Retraining Service
After=network.target

[Service]
User=ubuntu
WorkingDirectory=/home/ubuntu/ML_Model/ml_service
ExecStart=/home/ubuntu/ML_Model/ml_service/venv/bin/uvicorn api.app:app --host 0.0.0.0 --port 8000
Restart=always
RestartSec=5

[Install]
WantedBy=multi-user.target""")

    doc.add_paragraph("Enable and start the systemd background service:")
    add_code_block("""sudo systemctl daemon-reload
sudo systemctl enable ml_service
sudo systemctl start ml_service

# Check service status (should display active (running))
sudo systemctl status ml_service""")

    # 4. Docker Containerization Option
    add_heading_1("4. Docker Containerization Deployment (Alternative)")
    doc.add_paragraph("If your organization uses Docker or Kubernetes for deployment, use the built-in Dockerfile:")
    add_code_block("""# Navigate to ml_service root directory
cd ML_Model/ml_service

# Build Docker image
docker build -t nifty50-ml-service .

# Run Docker container with environment file and volume persistence
docker run -d \\
  --name nifty50_ml_app \\
  --restart always \\
  -p 8000:8000 \\
  --env-file .env \\
  -v $(pwd)/models/registry:/app/models/registry \\
  -v $(pwd)/data/live:/app/data/live \\
  nifty50-ml-service""")

    # 5. Reverse Proxy & HTTPS Domain Setup
    add_heading_1("5. Nginx Reverse Proxy & HTTPS Domain Setup")
    doc.add_paragraph("To allow web browsers to securely call your backend over HTTPS (https://api.yourdomain.com), setup Nginx and Let's Encrypt SSL:")
    add_code_block("""# Install Nginx and Certbot
sudo apt install -y nginx certbot python3-certbot-nginx

# Create Nginx site configuration
sudo nano /etc/nginx/sites-available/ml_service

# Add Nginx server block:
server {
    server_name api.yourdomain.com;

    location / {
        proxy_pass http://127.0.0.1:8000;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;
    }
}

# Enable site and issue free SSL certificate
sudo ln -s /etc/nginx/sites-available/ml_service /etc/nginx/sites-enabled/
sudo nginx -t
sudo systemctl restart nginx
sudo certbot --nginx -d api.yourdomain.com""")

    # 6. Connecting Frontend Web Application
    add_heading_1("6. Connecting the Frontend Web Application")
    doc.add_paragraph("Once the frontend team finishes building the website (React / Next.js / Vue), they connect to your deployed ML backend by configuring the API URL in their environment file:")

    add_heading_2("Step 6.1: Frontend Environment Configuration (.env.production)")
    add_code_block("""# Frontend .env.production
NEXT_PUBLIC_ML_API_URL=https://api.yourdomain.com""")

    add_heading_2("Step 6.2: API Endpoints Summary for Frontend Team")
    
    # API Table
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["HTTP Method & Endpoint", "Purpose", "Sample Usage"]
    for idx, h in enumerate(headers):
        cell = table.cell(0, idx)
        cell.paragraphs[0].text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)

    api_rows = [
        ("GET /predict/{ticker}", "Returns 30-min prediction (direction, return %, target price, confidence)", "fetch('/predict/EICHERMOT')"),
        ("GET /predict?tickers=...", "Batch predictions for multiple tickers for scanner UI", "fetch('/predict?tickers=WIPRO,TCS,RELIANCE')"),
        ("GET /health", "Returns system status, Fyers status & champion versions", "fetch('/health')"),
        ("GET /fyers/login", "Redirects user to Fyers 1-click OAuth login screen", "window.location.href = '/fyers/login'"),
        ("POST /retrain", "Triggers manual self-retraining job in background", "fetch('/retrain', {method: 'POST'})"),
    ]

    for row_idx, data in enumerate(api_rows, start=1):
        for col_idx, text in enumerate(data):
            cell = table.cell(row_idx, col_idx)
            cell.paragraphs[0].text = text
            if row_idx % 2 == 0:
                set_cell_background(cell, "F1F5F9")
            else:
                set_cell_background(cell, "FFFFFF")
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_heading_2("Step 6.3: Deploying Frontend to Vercel / Netlify")
    add_bullet("Vercel Deployment:", "Connect GitHub frontend repository to Vercel. Set Environment Variable NEXT_PUBLIC_ML_API_URL = https://api.yourdomain.com. Click Deploy.")
    add_bullet("Netlify Deployment:", "Connect GitHub repository, set build command 'npm run build', set environment variable, and deploy.")

    # 7. Automated Maintenance & Fyers Token Refresh
    add_heading_1("7. Automated Maintenance & Fyers Token Management")
    add_bullet("Self-Retraining Schedule:", "The background scheduler (retraining/scheduler.py) automatically triggers retrain_job.py every weekday at 6:00 PM (after market close). It pulls updated Fyers candles, retrains candidate models, evaluates the Champion Promotion Gate, and auto-promotes better models without server restarts.")
    add_bullet("Fyers 24H Token Renewal:", "Fyers access tokens expire after 24 hours. Users can click 'Connect Fyers Live' on the dashboard (/fyers/login) or call GET /fyers/login to renew the token in one click.")
    add_bullet("Drift Logs & Monitoring:", "Realized outcome errors are continuously logged to data/processed/drift_log.csv. If accuracy drops below 52%, an out-of-schedule emergency retrain is triggered.")

    # 8. Post-Deployment Verification Checklist
    add_heading_1("8. Post-Deployment Verification Checklist")
    doc.add_paragraph("Run through these sanity checks after completing deployment:")
    add_bullet("Check 1:", "Curl health endpoint: curl -s https://api.yourdomain.com/health | jq . (verify status: healthy).")
    add_bullet("Check 2:", "Test live prediction: curl -s https://api.yourdomain.com/predict/WIPRO | jq . (verify return %, target price, direction).")
    add_bullet("Check 3:", "Test alias resolution: curl -s https://api.yourdomain.com/predict/Tata%20Steel (verify mapped to TATASTEEL).")
    add_bullet("Check 4:", "Test ticker boundary guard: curl -s https://api.yourdomain.com/predict/ZOMATO (verify HTTP 400 rejection).")
    add_bullet("Check 5:", "Open frontend dashboard website in browser and verify live predictions render on UI cards.")

    # Save Word Document
    out_file = BASE_DIR.parent / "NIFTY50_ML_Deployment_Guide.docx"
    doc.save(out_file)
    print("Deployment guide successfully generated:", out_file)
    return out_file

if __name__ == "__main__":
    create_deployment_doc()
